# imports
import os
import json
import docx
import spotipy
import datetime
import spotipy.util as util
import spotipy.oauth2 as oauth2

from django.views import View
from django.db.models import Q
from django.conf import settings
from django.contrib import messages
from django.http import JsonResponse
from django.contrib.auth import get_user_model
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required, permission_required
from django.shortcuts import render, get_object_or_404, redirect, HttpResponse

from songs import models as song_models
from videos import models as video_models
from courses import forms as course_forms
from courses import models as course_models
from accounts import models as account_models

User = get_user_model()

# End: imports -----------------------------------------------------------------


# Functions
def trace(x):
    print(json.dumps(x, indent=4, sort_keys=True))


# End: Functions ---------------------------------------------------------------

addCourse_dec = [login_required, permission_required('courses.add_course', login_url='forbidden')]


@method_decorator(addCourse_dec, name='dispatch')
class AddCourseView(View):
    template = 'courses/course_form.html'
    courseForm_class = course_forms.CourseForm
    sectionForm_class = course_forms.SectionForm

    def get(self, request):
        courseForm = self.courseForm_class()
        sectionForm = self.sectionForm_class()
        sectionFormTemplate = self.sectionForm_class(prefix="template")
        return render(request, 'courses/course_form.html', {
            'courseForm': courseForm,
            'sectionForms': [],
            'sectionFormTemplate': sectionFormTemplate,
        })

    def post(self, request):

        courseForm = self.courseForm_class(data=request.POST)
        sectionCount = int(request.POST.get("sectionCount", "0"))

        sectionForms = []

        prefixes = request.POST.getlist("prefix")
        sectionForms = [self.sectionForm_class(prefix=prefixes[i], data=request.POST) for i in range(sectionCount)]

        if courseForm.is_valid():

            if all(sectionForm.is_valid() for sectionForm in sectionForms):
                course = courseForm.save()
                course.last_editor = request.user
                course.save()
                duration = 0
                # Add current sections
                for i in range(len(sectionForms)):
                    section = sectionForms[i].save()
                    section.nr = i + 1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.duration
                    section.save()

                return redirect('courses:course_view', courseID=course.id)

        return render(
            request, 'courses/course_form.html', {
                'courseForm': courseForm,
                'sectionForms': sectionForms,
                'sectionFormTemplate': self.sectionForm_class(prefix="template"),
            }
        )


editCourse_dec = [login_required, permission_required('courses.change_course', login_url='forbidden')]


@method_decorator(editCourse_dec, name='dispatch')
class EditCourseView(View):
    template = 'courses/course_form.html'
    courseForm_class = course_forms.CourseForm
    sectionForm_class = course_forms.SectionForm

    def get(self, request, courseID):
        course = course_models.Course.objects.get(id=courseID)
        sections = list(course.sections.all())
        courseForm = self.courseForm_class(instance=course)
        sectionForms = [self.sectionForm_class(prefix=i + 1, instance=sections[i]) for i in range(len(sections))]

        return render(
            request, self.template, {
                'courseForm': courseForm,
                'sectionForms': sectionForms,
                'courseID': course.id,
                'sectionFormTemplate': self.sectionForm_class(prefix="template"),
            }
        )

    def post(self, request, courseID):
        course = course_models.Course.objects.get(id=courseID)
        sections = list(course.sections.all())
        sectionForms = [self.sectionForm_class(prefix=i + 1, instance=sections[i]) for i in range(len(sections))]

        courseForm = self.courseForm_class(request.POST, instance=course)
        sectionCount = int(request.POST.get("sectionCount", "0"))

        prefixes = request.POST.getlist("prefix")
        sectionForms = [self.sectionForm_class(prefix=prefixes[i], data=request.POST) for i in range(sectionCount)]

        if courseForm.is_valid():

            if all(sectionForm.is_valid() for sectionForm in sectionForms):
                course = courseForm.save()
                course.last_editor = request.user
                course.save()
                course.sections.all().delete()  # reset sections

                duration = 0
                # Add current sections
                for i in range(len(sectionForms)):
                    section = sectionForms[i].save()
                    section.nr = i + 1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.duration
                    section.save()

                return redirect('courses:course_view', courseID=courseID)

        return render(
            request, self.template, {
                'courseForm': courseForm,
                'courseID': course.id,
                'sectionForms': sectionForms,
                'sectionFormTemplate': self.sectionForm_class(prefix="template"),
            }
        )


allCourses_dec = [login_required, permission_required('courses.view_course', login_url='forbidden')]


@method_decorator(allCourses_dec, name='dispatch')
class AllCoursesView(View):
    template = 'courses/all_courses.html'
    form = course_forms.CourseFilterForm

    def get(self, request):
        courses = course_models.Course.objects.all()
        form = self.form()
        return render(request, self.template, {'form': form, 'courses': courses})

    def post(self, request):
        courses = course_models.Course.objects.all()
        form = self.form(data=request.POST)
        if form.is_valid():
            courses = self.course_filter(form, courses)
        return render(request, self.template, {'form': form, 'courses': courses})

    def course_filter(self, form, queryset):
        search = form.cleaned_data['search']
        tag = form.cleaned_data['tag'] or None
        lead = form.cleaned_data['lead'] or None
        follow = form.cleaned_data['follow'] or None
        bulk = form.cleaned_data['bulk']
        day = form.cleaned_data['day']
        semester_char = form.cleaned_data['semester_char']
        external = form.cleaned_data['external']

        if search != "":
            queryset = queryset.filter(title__icontains=search)
        if tag:
            queryset = queryset.filter(tags__id=tag)
        if lead:
            queryset = queryset.filter(lead=lead)
        if follow:
            queryset = queryset.filter(follow=follow)
        if bulk:
            queryset = queryset.filter(bulk=bulk)
        if day:
            queryset = queryset.filter(day=day)
        if semester_char:
            queryset = queryset.filter(semester_char=semester_char)

        queryset = queryset.filter(external=external)

        return queryset


course_dec = [login_required, permission_required('courses.view_course', login_url='forbidden')]


@method_decorator(course_dec, name='dispatch')
class CourseView(View):
    template = 'courses/course_view.html'

    def get(
        self,
        request,
        courseID,
    ):
        course = course_models.Course.objects.get(id=courseID)
        return render(request, self.template, {'course': course})


deleteCourse_dec = [login_required, permission_required('courses.delete_course', login_url='forbidden')]


@method_decorator(deleteCourse_dec, name='dispatch')
class DeleteCourseView(View):

    def post(self, request, courseID):
        course_models.Course.objects.get(id=courseID).delete()
        messages.success(request, 'Course was successfully deleted')
        return redirect('courses:all_courses')


duplicate_dec = [login_required, permission_required('courses.create_course', login_url='forbidden')]


@method_decorator(duplicate_dec, name='dispatch')
class DuplicateCourse(View):

    def post(self, request, courseID):
        duplicate = course_models.Course.objects.get(id=courseID)
        duplicate.pk = None
        duplicate.title = duplicate.title + " (kopi)"
        duplicate.save()
        messages.success(request, 'Kurset har blitt duplisert')
        return redirect('courses:course_view', duplicate.id)


genPlaylist_dec = [login_required, permission_required('courses.view_course', login_url='forbidden')]


@method_decorator(genPlaylist_dec, name='dispatch')
class CreatePlaylistView(View):

    def get(self, request, courseID):
        course = course_models.Course.objects.get(id=courseID)

        # Auth client
        sp_oauth = oauth2.SpotifyOAuth(settings.SPOTIFY_CLIENT_ID, settings.SPOTIFY_CLIENT_SECRET, settings.SPOTIFY_REDIRECT_URI, scope=settings.SPOTIFY_SCOPE)

        sp_token, created = account_models.SpotifyToken.objects.get_or_create(user=request.user)

        token_info = None
        if sp_token.info:
            token_info = json.loads(sp_token.info)

            if sp_oauth._is_token_expired(token_info):
                token_info = sp_oauth.refresh_access_token(token_info['refresh_token'])
                sp_token.add_info(token_info)

            # if scopes don't match, then bail
            # if 'scope' not in token_info or not sp_oauth._is_scope_subset(sp_oauth.scope, token_info['scope']):
            #     messages.error(request, "scope mismatch")
            #     return redirect('courses:course_view', courseID=courseID)

            #return self.token_info['access_token']

        if not token_info:
            auth_url = sp_oauth.get_authorize_url()
            messages.error(request, 'Error. Dette skjer første gang du kobler til Spotify. Prøv igjen.')
            return redirect(auth_url + '&show_dialog=true')

        token = token_info['access_token']

        # Spotify API object
        spotify = spotipy.Spotify(auth=token)

        # Get username
        spotify_username = spotify.current_user()['uri'].split(':')[-1]

        # Title of playlist to be created
        playlist_title = '{} ({})'.format(course.title, course.date)

        # Get existing playlists for users
        playlists = spotify.user_playlists(spotify_username)

        # Check if playlist already exists
        playlist = [v for v in playlists['items'] if v['name'] == playlist_title]

        if not playlist:
            # Create new playlist
            playlist = spotify.user_playlist_create(user=spotify_username, name=playlist_title)
        else:
            # Take first mathing playlist (should be only one anyway)
            playlist = playlist[0]

        # Get URI for all songs used in course
        tracks = [section.song.spotify_URI for section in course.sections.all() if section.song]

        # Add tracks to playlist
        spotify.user_playlist_replace_tracks(user=spotify_username, playlist_id=playlist['id'], tracks=tracks)

        messages.success(request, f"{playlist_title} har blitt laget på Spotify kontoen din")
        return redirect('courses:course_view', courseID=courseID)


export_dec = [login_required, permission_required('courses.view_course', login_url='forbidden')]


@method_decorator(export_dec, name='dispatch')
class ExportView(View):

    def get(self, request, courseID):

        course = course_models.Course.objects.get(id=courseID)

        document = docx.Document()

        document.add_heading(course.title, level=1)

        tag_names = ", ".join(course.getTags())

        if course.lead:
            fører_navn = course.lead.get_full_name()
        else:
            fører_navn = 'Ingen instruktør valgt'

        if course.follow:
            følger_navn = course.follow.get_full_name()
        else:
            følger_navn = 'Ingen instruktør valgt'

        informasjon = "Instruktør (fører): {}\nInstruktør (følger): {}\nDato: {}\nNår: {} - {}\nHvor: {}\nTema: {}".format(
            fører_navn, følger_navn, course.getDate(), course.getStart(), course.getEnd(), course.place, tag_names
        )
        p = document.add_paragraph(informasjon)

        for section in course.sections.all():
            # p = document.add_paragraph("")
            h = document.add_heading("{} ({} min) - {}".format(section.title, section.duration, section.getStart()), level=2)

            p = document.add_paragraph()
            run = p.add_run("Sang: {}".format(section.getSong()))
            run.italic = True
            run.font.size = docx.shared.Pt(9)

            run = p.add_run("\n\n{}".format(section.description))

        h = document.add_heading("Kommentarer: ", level=2)
        p = document.add_paragraph(course.comments)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={}.docx'.format(course)
        document.save(response)

        return response

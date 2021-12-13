# imports
import json
import datetime

import docx
import spotipy

from django.conf import settings
from django.http import HttpRequest
from django.views import View
from django.contrib import messages
from django.shortcuts import render, redirect, HttpResponse
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required, permission_required

from accounts import models as account_models

from ..forms.courses import SectionForm, CourseFilterForm, CourseForm
from ..models.courses import Course, CourseSection

# End: imports -----------------------------------------------------------------


@method_decorator([login_required, permission_required('courses.add_course', login_url='forbidden')], name='dispatch')
class AddCourseView(View):
    template = 'ntnuisf/courses/course_form.html'
    course_form_class: CourseForm = CourseForm
    section_form_class: SectionForm = SectionForm

    def get(self, request: HttpRequest):
        course_form: CourseForm = self.course_form_class()
        # section_form = self.section_form_class()
        section_form_template = self.section_form_class(prefix='template')
        return render(request, 'courses/course_form.html', {
            'course_form': course_form,
            'section_forms': [],
            'section_form_template': section_form_template,
        })

    def post(self, request: HttpRequest):

        course_form = self.course_form_class(data=request.POST)
        section_count = int(request.POST.get('sectionCount', '0'))

        section_forms = []

        prefixes = request.POST.getlist('prefix')
        section_forms = [self.section_form_class(prefix=prefixes[i], data=request.POST) for i in range(section_count)]

        if course_form.is_valid():

            if all(section_form.is_valid() for section_form in section_forms):
                course = course_form.save()
                course.last_editor = request.user
                course.save()
                duration = 0
                # Add current sections
                for i, _ in enumerate(section_forms):
                    section = section_forms[i].save()
                    section.nr = i + 1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.duration
                    section.save()

                return redirect('courses:course_view', course_id=course.id)

        return render(
            request, 'courses/course_form.html', {
                'course_form': course_form,
                'section_forms': section_forms,
                'section_form_template': self.section_form_class(prefix='template'),
            }
        )


@method_decorator([login_required, permission_required('courses.change_course', login_url='forbidden')], name='dispatch')
class EditCourseView(View):
    template = 'ntnuisf/courses/course_form.html'
    course_form_class: CourseForm = CourseForm
    section_form_class: SectionForm = SectionForm

    def get(self, request: HttpRequest, course_id: int):
        course = Course.objects.get(id=course_id)
        sections = list(course.sections.all())
        course_form: CourseForm = self.course_form_class(instance=course)
        section_forms = [self.section_form_class(prefix=i + 1, instance=sections[i]) for i in range(len(sections))]

        return render(
            request, self.template, {
                'course_form': course_form,
                'section_forms': section_forms,
                'course_id': course.id,
                'section_form_template': self.section_form_class(prefix='template'),
            }
        )

    def post(self, request: HttpRequest, course_id: int):
        course = Course.objects.get(id=course_id)
        sections = list(course.sections.all())
        section_forms = [self.section_form_class(prefix=i + 1, instance=sections[i]) for i in range(len(sections))]

        course_form = self.course_form_class(request.POST, instance=course)
        section_count = int(request.POST.get('sectionCount', '0'))

        prefixes = request.POST.getlist('prefix')
        section_forms = [self.section_form_class(prefix=prefixes[i], data=request.POST) for i in range(section_count)]

        if course_form.is_valid():

            if all(section_form.is_valid() for section_form in section_forms):
                course = course_form.save()
                course.last_editor = request.user
                course.save()
                course.sections.all().delete()  # reset sections

                duration = 0
                # Add current sections
                for i, _ in enumerate(section_forms):
                    section = section_forms[i].save()
                    section.nr = i + 1
                    section.course = course
                    if course.start:
                        section.start = course.start + datetime.timedelta(minutes=duration)
                        duration += section.duration
                    section.save()

                return redirect('courses:course_view', course_id=course_id)

        return render(
            request, self.template, {
                'course_form': course_form,
                'course_id': course.id,
                'section_forms': section_forms,
                'section_form_template': self.section_form_class(prefix='template'),
            }
        )


allCourses_dec = [login_required, permission_required('courses.view_course', login_url='forbidden')]


@method_decorator(allCourses_dec, name='dispatch')
class AllCoursesView(View):
    template = 'ntnuisf/courses/all_courses.html'
    form: CourseFilterForm = CourseFilterForm

    def get(self, request: HttpRequest):
        courses = Course.objects.all()
        form: CourseFilterForm = self.form()
        return render(request, self.template, {'form': form, 'courses': courses})

    def post(self, request: HttpRequest):
        courses = Course.objects.all()
        form: CourseFilterForm = self.form(data=request.POST)
        if form.is_valid():
            courses = self.course_filter(form, courses)
        return render(request, self.template, {'form': form, 'courses': courses})

    def course_filter(self, form, queryset):  # pylint: disable=no-self-use
        search = form.cleaned_data['search']
        tag = form.cleaned_data['tag'] or None
        lead = form.cleaned_data['lead'] or None
        follow = form.cleaned_data['follow'] or None
        bulk = form.cleaned_data['bulk']
        day = form.cleaned_data['day']
        semester_char = form.cleaned_data['semester_char']
        external = form.cleaned_data['external']

        if search != '':
            queryset = queryset.filter(title__icontains=search)
        if tag:
            queryset = queryset.filter(tags__id=tag)
        if lead:
            queryset = queryset.filter(lead=lead)
        if follow:
            queryset = queryset.filter(follow=follow)
        if bulk:
            queryset = queryset.filter(bulk=bulk)
        if day:
            queryset = queryset.filter(day=day)
        if semester_char:
            queryset = queryset.filter(semester_char=semester_char)

        queryset = queryset.filter(external=external)

        return queryset


@method_decorator([login_required, permission_required('courses.view_course', login_url='forbidden')], name='dispatch')
class CourseView(View):
    template = 'ntnuisf/courses/course_view.html'

    def get(self, request: HttpRequest, course_id: int):
        course: Course = Course.objects.get(id=course_id)
        return render(request, self.template, {'course': course})


@method_decorator([login_required, permission_required('courses.delete_course', login_url='forbidden')], name='dispatch')
class DeleteCourseView(View):

    def post(self, request: HttpRequest, course_id: int):
        Course.objects.get(id=course_id).delete()
        messages.success(request, 'Course was successfully deleted')
        return redirect('courses:all_courses')


@method_decorator([login_required, permission_required('courses.create_course', login_url='forbidden')], name='dispatch')
class DuplicateCourse(View):

    def post(self, request: HttpRequest, course_id: int):
        duplicate_course: Course = Course.objects.get(id=course_id)
        duplicate_course.pk = None
        duplicate_course.title = duplicate_course.title + ' (kopi)'
        duplicate_course.save()
        messages.success(request, 'Kurset har blitt duplisert')
        return redirect('courses:course_view', duplicate_course.id)


@method_decorator([login_required, permission_required('courses.view_course', login_url='forbidden')], name='dispatch')
class CreatePlaylistView(View):

    def get(self, request: HttpRequest, course_id: int):
        course = Course.objects.get(id=course_id)

        # Auth client
        sp_oauth = spotipy.SpotifyOAuth(settings.SPOTIFY_CLIENT_ID, settings.SPOTIFY_CLIENT_SECRET, settings.SPOTIFY_REDIRECT_URI, scope=settings.SPOTIFY_SCOPE)

        sp_token, _created = account_models.SpotifyToken.objects.get_or_create(user=request.user)

        token_info = None
        if sp_token.info:
            token_info = json.loads(sp_token.info)

            if spotipy.SpotifyOAuth.is_token_expired(token_info):
                token_info = sp_oauth.refresh_access_token(token_info['refresh_token'])
                sp_token.add_info(token_info)

            # if scopes don't match, then bail
            # if 'scope' not in token_info or not sp_oauth._is_scope_subset(sp_oauth.scope, token_info['scope']):
            #     messages.error(request, 'scope mismatch')
            #     return redirect('courses:course_view', course_id=course_id)

            #return self.token_info['access_token']

        if not token_info:
            auth_url = sp_oauth.get_authorize_url()
            messages.error(request, 'Error. Dette skjer første gang du kobler til Spotify. Prøv igjen.')
            return redirect(auth_url + '&show_dialog=true')

        token = token_info['access_token']

        # Spotify API object
        spotify = spotipy.Spotify(auth=token)

        # Get username
        spotify_username = spotify.current_user()['uri'].split(':')[-1]

        # Title of playlist to be created
        playlist_title = f'{course.title} ({course.date})'

        # Get existing playlists for users
        playlists = spotify.user_playlists(spotify_username)

        # Check if playlist already exists
        playlist = [v for v in playlists['items'] if v['name'] == playlist_title]

        if not playlist:
            # Create new playlist
            playlist = spotify.user_playlist_create(user=spotify_username, name=playlist_title)
        else:
            # Take first mathing playlist (should be only one anyway)
            playlist = playlist[0]

        # Get URI for all songs used in course
        tracks = [section.song.spotify_URI for section in course.sections.all() if section.song]

        # Add tracks to playlist
        spotify.user_playlist_replace_tracks(user=spotify_username, playlist_id=playlist['id'], tracks=tracks)

        messages.success(request, f'{playlist_title} har blitt laget på Spotify kontoen din')
        return redirect('courses:course_view', course_id=course_id)


@method_decorator([login_required, permission_required('courses.view_course', login_url='forbidden')], name='dispatch')
class ExportView(View):

    def get(self, request: HttpRequest, course_id: int):

        course: Course = Course.objects.get(id=course_id)

        document = docx.Document()

        document.add_heading(course.title, level=1)

        tag_names = ', '.join(course.get_tags())

        if course.lead:
            lead_name = course.lead.get_full_name()
        else:
            lead_name = 'Ingen instruktør valgt'

        if course.follow:
            follow_name = course.follow.get_full_name()
        else:
            follow_name = 'Ingen instruktør valgt'

        info = f'Instruktør (fører): {lead_name}\nInstruktør (følger): {follow_name}\n \
            Dato: {course.get_date()}\nNår: {course.get_start()} - {course.get_end()}\nHvor: {course.place}\nTema: {tag_names}'

        p = document.add_paragraph(info)

        section: CourseSection
        for section in course.sections.all():
            # p = document.add_paragraph('')
            _ = document.add_heading(f'{section.title} ({section.duration} min) - {section.get_start()}', level=2)

            p = document.add_paragraph()
            run = p.add_run(f'Sang: {section.get_song()}')
            run.italic = True
            run.font.size = docx.shared.Pt(9)

            run = p.add_run(f'\n\n{section.description}')

        _ = document.add_heading('Kommentarer: ', level=2)
        p = document.add_paragraph(course.comments)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={course}.docx'
        document.save(response)

        return response
